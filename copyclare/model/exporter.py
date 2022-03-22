from copyclare.model.database import Database
from docx import Document
SAVE_AS = "../data/results"


class Exporter():
    def __init__(self, database):
        self.database = database
        self.doc = DocumentWriter()

    def export(self, report_title=None, attempt_id=None):
        quantitative_data = []
        qualitative_data = []
        export_title = ""
        if attempt_id is None:
            export_title = "Results for all attempts"
            attempts = self.database.get_all_attempts()
            for attempt in attempts:
                self.get_data(attempt, quantitative_data, qualitative_data)
        else:
            export_title = "Results for attempt %s" % attempt_id
            attempt = self.database.get_one_attempt_by_ID(attempt_id)
            self.get_data(attempt, quantitative_data, qualitative_data)
        self.doc.create_document(SAVE_AS + report_title,
                                 export_title, quantitative_data, qualitative_data)

    def get_data(self, attempt, quantitative_data, qualitative_data):
        exe_id = attempt.exercise_id
        exe = self.database.get_one_exercise_by_ID(exe_id)
        quantitative_data.append(
            {"name": exe.name,
             "reps": attempt.num_of_repetitons,
             "duration": attempt.duration,
             })
        qualitative_data.append(
            {"name": exe.name,
             "accuracy": attempt.accuracy,
             })

# reference: https://roytuts.com/a-guide-to-write-word-file-using-python/


class DocumentWriter():
    def __init__(self):
        self.document = Document()

    def create_document(self, saveAs, export_title, quantitative_data, qualitative_data):
        self.document.add_heading(export_title, 0)
        self.add_quantitative_section(quantitative_data)
        self.add_qualitative_section(qualitative_data)
        self.document.save(saveAs)

# which different exercises they performed
# number of repetition, time take
    def add_quantitative_section(self, quantitative_data):
        self.document.add_heading('Quantitative', level=1)
        table = self.document.add_table(rows=1, cols=3)
        hdr_cells = table.rows[0].cells
        hdr_cells[0].text = 'Exercise Name'
        hdr_cells[1].text = 'Repetitions'
        hdr_cells[2].text = 'Duration'
        for item in quantitative_data:
            row_cells = table.add_row().cells
            row_cells[0].text = str(item['name'])
            row_cells[1].text = str(item['reps'])
            row_cells[2].text = str(item['duration'])


# What we would ideally like to see is information
# about the person as they made the movement,
# and how that movement matched against the
# normal movement/reference movement.
# Ideally, we would like to get whole limb through
# range read out(Basically, how much they deviated
# from the normal range). For this we would want whole
# limb through range data in relation to the reference movement.

    def add_qualitative_section(self, qualitative_data):
        self.document.add_heading('Qualitative', level=1)
        table = self.document.add_table(rows=1, cols=2)
        hdr_cells = table.rows[0].cells
        hdr_cells[0].text = 'Exercise Name'
        hdr_cells[1].text = 'Accuracy'
        for item in qualitative_data:
            row_cells = table.add_row().cells
            row_cells[0].text = str(item['name'])
            row_cells[1].text = str(item['accuracy'])
