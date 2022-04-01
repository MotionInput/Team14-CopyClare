"""
Contributors: Yan Lai, Sree Sanakkayala

"""

"""

This export module deals with everything related to exporting progress reports for the users

.. code-block:: python

   # sample usage:
   from copyclare.data import Exporter
   # Exporter is an object that allows for the exporting of progress reports
"""


import os
import json
import pyqtgraph as pg
from pyqtgraph import exporters
from docx import Document
from docx.shared import Inches
from copyclare.data import DATA_DIR
DOCX_MAX_IMAGE_WIDTH = Inches(6.5)


class Exporter():
    """
    This class has the functions to gather all of the necessary information
    to from the given database to create a progress report.

    Args:
        database : The database from which that data will be exported from
    """

    def __init__(self, database):
        self.database = database
        self.doc = DocumentWriter()

    def export(self, saveAs, attempt_id=None):
        """
        Exports the data from the database to a .docx file.
        Args:
            saveAs : absolute path for where report needs to be saved.
            attempt_id : the attempt id of the attempt that needs to be exported.
                         if null, then the tool will export all

        """
        exercise_info = []
        quantitative_data = []
        qualitative_data = []
        export_title = ""
        if attempt_id is None:
            export_title = "Results for all attempts"
            attempts = self.database.get_all_attempts()
            for attempt in attempts:
                self._get_data(
                    attempt, exercise_info, quantitative_data, qualitative_data)
        else:
            export_title = "Results for attempt %s" % attempt_id
            attempt = self.database.get_one_attempt_by_ID(attempt_id)
            self._get_data(attempt, exercise_info, quantitative_data,
                           qualitative_data)
        self.doc.create_document(saveAs, export_title, exercise_info,
                                 quantitative_data, qualitative_data)

    def _get_data(self, attempt, exercise_info, quantitative_data,
                  qualitative_data):
        """This is a helper function that gathers all of the data

        Args:
            attempt (_type_): The attempt that needs to be exported
            exercise_info (_type_): the basic info of the exercise
            quantitative_data (_type_): the quantitative data of the exercise (shown below)
            qualitative_data (_type_): the qualitative data of the exercise (shown below)
        """
        exe_id = attempt.exercise_id
        exe = self.database.get_one_exercise_by_ID(exe_id)
        exercise_info.append({
            "name": exe.name,
            "image": DATA_DIR + exe.image_directory,
            "description": exe.description,
            "date": attempt.date,
            "id": exe.id,
        })
        quantitative_data.append({
            "reps": attempt.num_of_repetitons,
            "duration": round(attempt.duration, 2),
        })
        qualitative_data.append({
            "accuracy": round(attempt.accuracy, 2),
            "accuracy_graph": DATA_DIR+f"/accuracy-graphs/{attempt.id}.png",
        })


class DocumentWriter():
    """ This is a helper class build the .docx file based on the inputted data
    """

    def __init__(self):
        self.document = Document()

    def create_document(self, saveAs, export_title, exercise_info,
                        quantitative_data, qualitative_data):
        """the main function to call to create a .docx file

        Args:
            saveAs (_type_): absolute path of where the file needs to be stored
            export_title (_type_): the title of the document
            exercise_info (_type_): basic info of the exercise
            quantitative_data (_type_): quantitative data of the exercise (see above)
            qualitative_data (_type_): qualitative data of the exercise (see above)
        """
        self.document.add_heading(export_title, 0)
        self._add_all_progress_charts(exercise_info)
        for i in range(len(exercise_info)):
            self._add_exercise_details(exercise_info[i])
            self._add_quantitative_section(quantitative_data[i])
            self._add_qualitative_section(qualitative_data[i])
        self.document.save(saveAs)

    def _add_all_progress_charts(self, exercise_info):
        """this is a helper to first add all the progress charts to the report

        Args:
            exercise_info (dict): basic info for the exercise
        """
        displayed_exercises = set()
        for exercise in exercise_info:
            if exercise["id"] not in displayed_exercises:
                self.document.add_heading('Progress Chart for %s over time' %
                                          exercise["name"], level=2)
                exe_id = exercise["id"]
                self.document.add_picture(
                    DATA_DIR+f"/progress-charts/{exe_id}.png", width=DOCX_MAX_IMAGE_WIDTH)
                displayed_exercises.add(exe_id)

    def _add_exercise_details(self, exercise_info):
        """this is a helper to add the name and description of the exercise

        Args:
            exercise_info (dict): the basic info of the exercise
        """
        self.document.add_heading('Exercise Name: %s - %s' % (exercise_info["name"], exercise_info["date"]),
                                  level=2)
        self.document.add_picture(
            exercise_info["image"], width=DOCX_MAX_IMAGE_WIDTH)
        self.document.add_paragraph('Description:  %s' %
                                    exercise_info["description"])

    def _add_quantitative_section(self, quantitative_data):
        """this is a helper to add the quantitative data of the exercise

        Args:
            quantitative_data (dict): the quantitative data of the exercise
        """
        self.document.add_heading('Quantitative', level=2)
        table = self.document.add_table(rows=1, cols=2)
        hdr_cells = table.rows[0].cells
        hdr_cells[0].text = 'Repetitions'
        hdr_cells[1].text = 'Duration'
        row_cells = table.add_row().cells
        row_cells[0].text = str(quantitative_data['reps'])
        row_cells[1].text = str(quantitative_data['duration']) + " Seconds"

    def _add_qualitative_section(self, qualitative_data):
        """this is a helper function to add qualitative_data to the report

        Args:
            qualitative_data (dict): qualitative data of the exercise
        """
        self.document.add_heading('Qualitative', level=2)
        table = self.document.add_table(rows=1, cols=1)
        hdr_cells = table.rows[0].cells
        hdr_cells[0].text = 'Avg. Accuracy'
        row_cells = table.add_row().cells
        row_cells[0].text = str(qualitative_data['accuracy'])+"%"
        self.document.add_heading('Accuracy graph for attempt', level=3)
        self.document.add_picture(
            qualitative_data["accuracy_graph"], width=DOCX_MAX_IMAGE_WIDTH)


class AccuracyGraphExporter:
    def draw_accuracy_graph(self, session_json):
        """
        Args:
            session_json ([[int,int]]]): A list of tuples containing timestamps and accuracy for an attempt.

        Returns:
            PlotWidget: The plotted graph widget.

        """
        graphWidget = pg.PlotWidget()

        x_axis = []
        y_axis = []

        accuracy = json.loads(session_json)

        for pair in accuracy:
            x_axis.append(pair[0])  # timestamp
            y_axis.append(pair[1])  # accuracy

        graphWidget.setBackground('w')
        graphWidget.setTitle(
            "Accuracy throughout exercise (recorded by second)", color="black", size="15pt")
        graphWidget.showGrid(x=True, y=True)
        if len(x_axis) == 0:
            graphWidget.setXRange(0, 0, padding=0)
        else:
            graphWidget.setXRange(x_axis[0], x_axis[-1], padding=0)
        graphWidget.setYRange(0, 100, padding=0)

        pen = pg.mkPen(color=(0, 20, 40), width=3)
        graphWidget.plot(x_axis, y_axis, name="",  pen=pen,
                         symbol='o', symbolSize=2, symbolBrush=('#003366'))

        return graphWidget

    def export_accuracy_graph(self, session_json, attempt_id):
        """
        Args:
            session_json ([[int,int]]]): A list of tuples containing timestamps and accuracy for an attempt.
            attempt_id (int): The id of selected attempt.


        """
        path = DATA_DIR + f'/accuracy-graphs/{attempt_id}.png'
        exists = os.path.exists(path)
        if not exists:
            accuracy_graph = self.draw_accuracy_graph(session_json)
            exporter = exporters.ImageExporter(accuracy_graph.plotItem)
            exporter.parameters()['width'] = 500
            exporter.parameters()['height'] = 400
            exporter.export(path)
